# -*- coding: utf-8 -*-
"""Generate the MiniChess AI report (Traditional Chinese) as .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = "1E2761"
BLUE = "3B5BA9"
TEXT = "1A2238"
MUTED = "5B6577"
GREEN = "2E7D4F"
RED = "C0392B"
HEADBG = "1E2761"

doc = Document()

# ---- page setup: A4, 2.2cm margins ----
sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Mm(22))

def set_fonts(run, latin="Microsoft JhengHei", ea="Microsoft JhengHei"):
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.append(rf)
    rf.set(qn("w:ascii"), latin)
    rf.set(qn("w:hAnsi"), latin)
    rf.set(qn("w:cs"), latin)
    rf.set(qn("w:eastAsia"), ea)

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(TEXT)
set_fonts_style = normal.element.get_or_add_rPr()
_rf = OxmlElement("w:rFonts")
for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
    _rf.set(qn(a), "Microsoft JhengHei")
set_fonts_style.append(_rf)

def style_heading(name, size, color):
    st = doc.styles[name]
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), "Microsoft JhengHei")

style_heading("Heading 1", 16, NAVY)
style_heading("Heading 2", 13, NAVY)
style_heading("Heading 3", 11.5, BLUE)

def add_runs(p, runs, size=11):
    for seg in runs:
        if isinstance(seg, str):
            seg = (seg,)
        txt = seg[0]
        bold = seg[1] if len(seg) > 1 else False
        color = seg[2] if len(seg) > 2 else None
        mono = seg[3] if len(seg) > 3 else False
        r = p.add_run(txt)
        r.font.size = Pt(9.5 if mono else size)
        r.font.bold = bold
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
        set_fonts(r, "Consolas" if mono else "Microsoft JhengHei", "Microsoft JhengHei")

def para(runs, size=11, after=6, align=None, indent=None, before=0):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    if isinstance(runs, str):
        runs = [runs]
    add_runs(p, runs, size)
    return p

def bullet(runs, size=11, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    if isinstance(runs, str):
        runs = [runs]
    add_runs(p, runs, size)
    return p

def H(text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    set_fonts(r)
    return h

def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)

def code_block(lines):
    p = doc.add_paragraph()
    shade(p, "F2F4F8")
    p.paragraph_format.left_indent = Pt(10)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for i, ln in enumerate(lines):
        r = p.add_run(ln)
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string("12306B")
        set_fonts(r, "Consolas", "Microsoft JhengHei")
        if i < len(lines) - 1:
            r.add_break()
    return p

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def add_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.allow_autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], HEADBG)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
        set_fonts(r)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            set_fonts(r)
    for i, w in enumerate(widths):
        for row in t.rows:
            row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ============================================================ Title
ttl = doc.add_paragraph()
ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ttl.add_run("MiniChess AI 演算法報告")
r.font.size = Pt(24)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(NAVY)
set_fonts(r)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Mini Project 2 — 搜尋演算法(Baseline)與優化(Algorithm)")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor.from_string(MUTED)
set_fonts(r)
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run("姓名/學號:＿＿＿＿＿＿　　日期:＿＿＿＿＿＿")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor.from_string(MUTED)
set_fonts(r)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# Overview box
ov = doc.add_paragraph()
shade(ov, "EEF2FB")
ov.paragraph_format.space_before = Pt(4)
ov.paragraph_format.space_after = Pt(6)
add_runs(ov, [
    ("一句話總結:", True, NAVY),
    ("本專案實作一個以 Negamax 為基礎、採用迭代加深的 ", False, TEXT),
    ("PVS(主變例搜尋)", True, NAVY),
    (" 引擎;以 alpha-beta + PVS 為核心,搭配置換表、走法排序、靜止搜尋、"
     "以及延伸/縮減/裁剪等選擇性深度技術,餵給一個結合材料、位置、王安全與兵形的評估函數。", False, TEXT),
])

# Contents
H("報告大綱", 2)
for t in ["1. 專案概述", "2. Baseline:核心搜尋與評估(對應 baseline 8 分)",
          "3. Algorithm:搜尋優化(對應 algorithm 4 分)",
          "4. 實驗結果", "5. TA 問答準備", "6. 附錄:關鍵參數"]:
    bullet(t)

# ============================================================ 1 Overview
H("1. 專案概述", 1)
H("1.1 遊戲規則重點", 2)
para([("棋盤為 6×5(6 列 5 行),雙方各有 王、后、車、象、馬 與 5 個兵。"
       "勝負規則直接影響評估與搜尋設計:", False, TEXT)])
bullet([("吃王即勝:", True, NAVY), ("只要有一手能吃掉對方的王(piece=6),立即獲勝"
        "(在 get_legal_actions() 中偵測,設定 game_state=WIN)。", False, TEXT)])
bullet([("將死:", True, NAVY), ("對方被將軍且所有應手仍會被吃王。", False, TEXT)])
bullet([("100 步子力判定:", True, NAVY), ("若 100 步內未吃王,以子力總和判勝;相等則和局。", False, TEXT)])
bullet([("和局:", True, NAVY), ("同一局面出現 3 次(以 Zobrist hash 偵測)。", False, TEXT)])
bullet([("判負:", True, NAVY), ("輪到一方卻無合法走法(且非可吃王狀態);或未在時限內輸出走法。", False, TEXT)])

H("1.2 程式架構與可編輯範圍", 2)
para([("可編輯檔案僅 ", False, TEXT), ("state.cpp / state.hpp", False, NAVY, True),
      (" 與 ", False, TEXT), ("src/policy/*", False, NAVY, True),
      (";框架檔 ", False, TEXT), ("ubgi.cpp", False, NAVY, True),
      ("(UBGI 通訊與時間控制)不可更動。搜尋與評估邏輯集中於:", False, TEXT)])
bullet([("src/policy/minimax.cpp", False, NAVY, True), ("　搜尋:MiniMax、AlphaBeta、PVS 三種演算法", False, TEXT)])
bullet([("src/games/minichess/state.cpp", False, NAVY, True), ("　評估函數、走法生成、Zobrist 雜湊", False, TEXT)])
bullet([("src/policy/registry.hpp", False, NAVY, True), ("　演算法註冊;預設名稱 \"minimax\" 實際對應到 PVS", False, TEXT)])

H("1.3 演算法總覽", 2)
para([("引擎提供三種搜尋,複雜度由淺到深堆疊,實際對局使用最強的 PVS:", False, TEXT)])
add_table(
    ["演算法", "技術", "用途"],
    [["MiniMax", "純 negamax,無剪枝", "Baseline 對照"],
     ["AlphaBeta", "negamax + alpha-beta + MVV-LVA", "Baseline 剪枝"],
     ["PVS", "PVS + 置換表 + 靜止搜尋 + 殺手 + LMR + 空著裁剪", "實際對局(優化版)"]],
    [1.4, 3.2, 1.7])

# ============================================================ 2 Baseline
H("2. Baseline:核心搜尋與評估", 1)
para([("(對應評分 baseline 8 分)", False, MUTED)])

H("2.1 盤面表示與走法生成", 2)
para([("盤面以 ", False, TEXT), ("board[2][6][5]", False, NAVY, True),
      (" 表示(兩方各一層)。子力編碼:", False, TEXT)])
add_table(["編碼", "1", "2", "3", "4", "5", "6"],
          [["棋子", "兵", "車", "馬", "象", "后", "王"]],
          [1.2, 0.88, 0.88, 0.88, 0.88, 0.88, 0.88])
bullet([("走法生成有兩版:", True, NAVY), ("陣列式(naive)與位元棋盤(bitboard,30 格放入 uint32_t,以 __builtin_ctz 掃描),"
        "結果相同但 bitboard 較快。", False, TEXT)])
bullet([("吃王偵測:", True, NAVY), ("生成走法時若發現可吃對方王,立即設定 game_state=WIN 並回傳。", False, TEXT)])
bullet([("兵的升變:", True, NAVY), ("兵走到底線(next_state 中)直接升為后(moved=5)。", False, TEXT)])

H("2.2 評估函數 evaluate()", 2)
para([("評估從「走子方」角度計算,", False, TEXT),
      ("分數 = 我方總分 − 對方總分 + 機動力加成", True, NAVY),
      (";正分代表對走子方有利。若為勝局直接回傳 P_MAX。組成:", False, TEXT)])
add_table(["項目", "說明"],
          [["材料 Material", "兵20 馬60 象65 車100 后200 王1000(比例符合官方 100 步判定)"],
           ["位置表 PST", "每種子在每格的位置加分(馬偏中央、王貼後排、兵越前越好)"],
           ["王接近度 Tropism", "我方子力靠近敵王給分(攻擊壓力);權重 {兵1 車4 馬4 象3 后6}"],
           ["通路兵 Passed Pawn", "前方無敵兵阻擋的兵,越接近升變加分越高 {60,30,14,4}"],
           ["車開放線", "車在全開放線 +12、半開放線 +6"],
           ["兵形 Pawn Structure", "疊兵 −15/個、孤兵 −10"],
           ["機動力 Mobility", "3 ×(我方可走步數 − 對方可走步數);用快速偽機動力計數"]],
          [1.7, 4.6])

H("2.3 Negamax 搜尋", 2)
para([("零和賽局下,以單一視角統一雙方:每個節點取最大,並對子節點取負。", False, TEXT)])
code_block(["score = -eval_ctx(child, depth-1, ...);   // negamax 核心",
            "勝局回傳  P_MAX - ply   // 減 ply 讓引擎偏好更快的將死"])
para([("P_MAX = 100000、M_MAX = −100000。終端條件:可吃王(WIN)、和局(0)、深度為 0(交給評估/靜止搜尋)。", False, TEXT)])

H("2.4 Alpha-Beta 剪枝", 2)
para([("維護窗口 [α, β]:α 是我方已保證的下限,β 是對手會容許的上限。"
       "當某手 score ≥ β,對手不會走到此局面,於是停止搜尋此節點(beta cutoff)。", False, TEXT)])
bullet([("結果與完整 minimax 相同", True, NAVY), (",但能剪掉大量分枝。", False, TEXT)])
bullet([("最佳排序下複雜度約 ", False, TEXT), ("O(b^(d/2))", False, NAVY, True),
        (",等於相同時間內深度可加倍。", False, TEXT)])
bullet([("搭配 ", False, TEXT), ("MVV-LVA", True, NAVY),
        (" 走法排序(吃最大子、用最小子吃)讓好棋先試,提高剪枝率。", False, TEXT)])

# ============================================================ 3 Algorithm
H("3. Algorithm:搜尋優化", 1)
para([("(對應評分 algorithm 4 分)", False, MUTED)])

H("3.1 PVS 主變例搜尋", 2)
para([("PVS 假設(排序良好時)第一手最好,用最省力方式驗證其餘手更差:", False, TEXT)])
bullet([("第一手用完整窗口 [α, β] 搜尋 → 建立主變例(PV)。", False, TEXT)])
bullet([("其餘手用零寬窗口 [α, α+1] 只做「是否更好」的布林測試,快很多。", False, TEXT)])
bullet([("若零寬測試意外 > α,才用完整窗口重新搜尋取得真值。", False, TEXT)])
para([("排序良好時重搜很少,因此 PVS 明顯快於一般 alpha-beta。", False, MUTED)])

H("3.2 迭代加深與時間控制", 2)
para([("不是一次搜固定深度,而是深度 1、2、3… 直到時間用完(迴圈位於 ubgi.cpp 的 do_search)。"
       "好處:隨時有完整最佳手可用;且每層用置換表餵給下一層,讓深層搜尋更快。", False, TEXT)])
bullet([("時間判斷:", True, NAVY), ("每完成一層檢查 total×2 ≥ movetime 就停;另外找到將死分數也提前結束。", False, TEXT)])
bullet([("時限機制:", True, NAVY), ("測試程式送 go movetime T 後睡 T 毫秒再 kill,取最後回報的 bestmove/currmove;"
        "故非「超時即判負」,而是取當下回報的最佳手。", False, TEXT)])

H("3.3 置換表(Transposition Table)", 2)
para([("同一局面常被不同走序到達。以 100 萬筆(TT_BITS=20)的雜湊表快取結果:", False, TEXT)])
bullet([("Zobrist 雜湊:", True, NAVY), ("每個(子,格,顏色)對應一隨機 64-bit 數,盤面雜湊=各子 XOR;"
        "走子時只需 XOR 出舊位置、XOR 入新位置即可 O(1) 增量更新。", False, TEXT)])
bullet([("內容:", True, NAVY), ("key(驗證碰撞)、分數、深度、界限類型(精確/下界/上界)、最佳走法。", False, TEXT)])
bullet([("深度優先替換", True, NAVY), (";將死分數依距根步數 ply 正規化(存入加 ply、取出減 ply),避免不同深度取用造成步數錯亂。", False, TEXT)])

H("3.4 走法排序", 2)
para([("剪枝效果取決於好棋是否先試。排序順序:", False, TEXT)])
para([("置換表走法", True, NAVY), ("  →  ", False, MUTED), ("吃子(MVV-LVA)", True, NAVY),
      ("  →  ", False, MUTED), ("殺手走法", True, NAVY), ("  →  ", False, MUTED),
      ("其餘安靜走法", True, NAVY)], after=4)
para([("殺手走法:每層保留 2 個曾造成 beta cutoff 的安靜走法,優先嘗試。", False, MUTED)])

H("3.5 靜止搜尋(Quiescence)", 2)
para([("解決「視界效應」——在深度 0 直接評估可能剛好停在被吃子之前而誤判佔優。", False, TEXT)])
bullet([("葉節點只繼續追「吃子」直到局面安靜才評估(上限 MAX_QDEPTH=8)。", False, TEXT)])
bullet([("stand-pat 為「不吃」的基準;", False, TEXT), ("delta pruning", True, NAVY),
        (" 跳過即使吃到也無法拉回 α 的吃子(餘裕 +50)。", False, TEXT)])

H("3.6 選擇性深度", 2)
add_table(["技術", "條件 / 作法"],
          [["將軍延伸 Check Ext.", "該手使對方被將軍(is_king_attacked)且 ply<20 → 多搜 1 層"],
           ["後期走法縮減 LMR", "depth≥3 且排序第 3 手後 且 非吃子非將軍 → 減 1(depth≥6 減 2);意外 >α 才重搜"],
           ["空著裁剪 Null-Move", "跳一手後仍 ≥ β 就剪枝;接近將死、低子力、被將軍時跳過(R=3,depth≥6 為 4)"],
           ["期望窗口 Aspiration", "以上一層分數 ±50 為窗口;失敗則 ×3 放寬重搜"]],
          [1.9, 4.4])

H("3.7 關鍵優化:根節點存入置換表(本專案最大修正)", 2)
p = doc.add_paragraph()
shade(p, "FBECEC")
p.paragraph_format.space_after = Pt(6)
add_runs(p, [("問題:", True, RED),
             ("原本 PVS::search 每一層都重新以吃子值(MVV-LVA)排序根節點走法,"
              "會把「吃大子但其實會輸的吃子」排在第一個。由於測試程式時間一到就 kill 引擎並取"
              "「最後回報的走法」,當搜尋在深層被中斷時,就會輸出那個敗著——導致開局第 7 手"
              "就被將死。固定深度搜尋本來都正確,只有「限時」這條路會出錯。", False, TEXT)])
p = doc.add_paragraph()
shade(p, "EAF5EE")
p.paragraph_format.space_after = Pt(6)
add_runs(p, [("解法:", True, GREEN),
             ("在 PVS::search 每層搜尋結束時,把根節點的最佳走法寫入置換表。如此下一層"
              "會先搜、先回報真正的最佳手;即使被時間中斷,回報的也是合理走法。", False, TEXT)])
para([("效果:對 Weak 由 ", False, TEXT), ("40% → 100%", True, GREEN),
      (",同一修正也讓 Strong、Boss 全面提升(見第 4 節)。", False, TEXT)])

# ============================================================ 4 Results
H("4. 實驗結果", 1)
para([("測試設定:每步 2000ms、各 10 局、雙方輪流先手(以專案附帶的對手執行檔對戰)。", False, TEXT)])
add_table(["對手", "修正前", "修正後"],
          [["Weak", "40%", "100%(10:0)"],
           ["Strong", "約 90%", "100%(10:0)"],
           ["Boss", "約 70%", "100%(10:0)"]],
          [2.1, 2.1, 2.1])
para([("結論:根節點存入置換表的修正,搭配既有的 PVS/置換表/排序/靜止搜尋等技術,"
       "讓引擎對三個對手皆達到穩定高勝率。", False, MUTED)])

# ============================================================ 5 Q&A
H("5. TA 問答準備", 1)
qa = [
    ("為什麼用 negamax 而不是分開寫 min/max?",
     "零和賽局下一方得分=另一方失分。negamax 讓每個節點都「取最大」並對子節點取負"
     "(score=−search(child)),用單一函式處理雙方,程式更簡潔、較不易出錯。"),
    ("alpha-beta 為何能在不改變結果下剪枝?最佳複雜度?",
     "α 是我方已保證的下限、β 是對手會容許的上限;當某手 score≥β,對手不會走到此局面,"
     "後續無需再算(beta cutoff),結果與完整 minimax 相同。最佳排序下約 O(b^(d/2)),深度可加倍。"),
    ("PVS 與 alpha-beta 差在哪?零寬窗口是什麼?",
     "PVS 先用完整窗口算出第一手(主變例),其餘手改用零寬窗口 [α,α+1] 只做「是否更好」的"
     "布林測試(快很多);若意外 >α 才用完整窗口重搜。排序良好時重搜少,因此比一般 alpha-beta 快。"),
    ("走法排序為何重要?你怎麼排?",
     "剪枝效果取決於好棋是否先試。順序為:置換表走法 → 吃子(MVV-LVA)→ 殺手走法 → 其餘安靜走法。"
     "好棋先試 → 更早 cutoff → 看更深。"),
    ("置換表如何處理碰撞與將死分數?",
     "entry 內存 key 做驗證,碰撞時 key 不符就不採用;採深度優先替換。將死分數與距根步數 ply 有關,"
     "存入時加 ply、取出時減 ply 做正規化,避免不同深度取用造成將死步數錯亂。"),
    ("靜止搜尋解決什麼?delta pruning 是什麼?",
     "解決「視界效應」:深度 0 直接評估可能停在被吃子之前而誤判。靜止搜尋只追吃子直到安靜才評估;"
     "delta pruning 跳過即使吃到也無法拉回 α 的吃子以省時。"),
    ("LMR 與空著裁剪有什麼風險?何時不用?",
     "兩者皆為啟發式,可能漏算戰術。LMR 只縮減「排序後段、非吃子、非將軍」的走法,且意外 >α 時"
     "以完整深度重搜;空著裁剪在接近將死、低子力(易 zugzwang)、自己被將軍時跳過。"),
    ("你最大的優化是什麼?(預期會被問)",
     "把『根節點最佳走法存入置換表』。原本每層用吃子排序會把會輸的吃子排第一,限時被中斷時就走出敗著;"
     "修正後每層先搜先報真正最佳手,對 Weak 由 40% 提升到 100%。"),
    ("時間如何控制?會不會超時判負?",
     "採迭代加深,每完成一層檢查 total×2≥movetime 就停。測試程式送 go movetime T 後睡 T 再 kill,"
     "取最後回報的走法;故非超時即判負,真正的失分來源是上述根節點排序問題(已修正)。"),
    ("評估函數各項意義?分數正負代表什麼?",
     "分數=我方總分−對方總分(走子方視角,正=有利)。含材料、位置表、王接近度、通路兵、"
     "車開放線、兵形(疊兵/孤兵扣分)、機動力。"),
    ("為什麼用 Zobrist hash 而非整個盤面當鍵?",
     "Zobrist 把每個(子,格,顏色)對應一隨機 64-bit 數,盤面雜湊=各子 XOR;走子時只需 XOR 出/入"
     "兩格即可 O(1) 增量更新,且分布均勻,非常適合當雜湊表鍵。"),
    ("100 步判定如何影響你的策略?",
     "若 100 步內未吃王,依子力多寡判勝負(官方:兵1 車5 馬3 象3 后9)。故長局需確保子力領先;"
     "引擎的材料權重比例與此一致。"),
]
for i, (q, a) in enumerate(qa, 1):
    para([("Q%d. " % i, True, NAVY), (q, True, NAVY)], after=2, before=4)
    para([("A:", True, BLUE), (" " + a, False, TEXT)], after=4, indent=12)

# ============================================================ 6 Appendix
H("6. 附錄:關鍵參數", 1)
add_table(["參數", "值 / 條件"],
          [["P_MAX / M_MAX", "+100000 / −100000(勝/負分數上下限)"],
           ["置換表大小", "TT_BITS=20 → 1,048,576 筆(約 24MB)"],
           ["靜止搜尋上限", "MAX_QDEPTH = 8"],
           ["材料(評估)", "兵20 馬60 象65 車100 后200 王1000"],
           ["材料(官方判定)", "兵1 車5 馬3 象3 后9 王0"],
           ["LMR", "depth≥3 且 第3手後 且 非吃子非將軍 → 減1(depth≥6 減2)"],
           ["空著裁剪", "depth≥3、beta<P_MAX−100、非兵子力≥2;R=3(depth≥6 為4)"],
           ["將軍延伸", "ply<20 且該手使對方被將軍 → +1 層"],
           ["期望窗口", "±50,失敗 ×3 放寬"],
           ["殺手走法", "每層 2 個"],
           ["兵形罰分", "疊兵 −15/個、孤兵 −10"]],
          [2.0, 4.3])

doc.save("MiniChess_AI_報告.docx")
print("saved MiniChess_AI_報告.docx")
